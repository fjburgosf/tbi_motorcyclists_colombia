# -*- coding: utf-8 -*-
"""
Build a conventional (non-MDPI) Word .docx from manuscript_safety_final_EN.md.

Features:
  * Body text with **bold**, *italic* and `code` inline formatting.
  * Headings, paragraphs, pipe tables and PNG figures embedded.
  * The statistical model equations (described in prose in the source) embedded
    as native Word (OMML) equation objects.
  * The significance-asterisk convention explained in the captions of the
    statistical tables.
"""

import re
import struct
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml

# --------------------------------------------------------------------------- paths
BASE = Path(r"D:/ACADEMICO/Papers/TBI_data_analysis")
MD_PATH = BASE / "manuscript" / "manuscript_safety_final_EN.md"
OUT_PATH = BASE / "manuscript" / "manuscript_safety_final_EN.docx"
FIG_DIR = BASE / "figures"

BODY_FONT = "Times New Roman"

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSDECL_M = 'xmlns:m="%s"' % M_NS
NSDECL_W = 'xmlns:w="%s"' % W_NS

# Significance legend appended to captions of tables that report p-values / models.
SIG_LEGEND = (
    " Significance convention (asterisk notation): *** p < 0.001, ** p < 0.01, "
    "* p < 0.05, n.s. = not significant. Exact p-values are reported."
)
SIG_TABLES = {"Table 2", "Table 3", "Table S3", "Table S4"}

# -------------------------------------------------------------------- OMML helpers
def _esc(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def m_r(t, sty=None):
    """A math run. sty=None -> italic (default math), 'p' -> upright/plain."""
    rpr = ('<m:rPr><m:sty m:val="%s"/></m:rPr>' % sty) if sty else ""
    return '<m:r>%s<m:t xml:space="preserve">%s</m:t></m:r>' % (rpr, _esc(t))


def m_txt(t):          # upright literal text (operators, words, digits, parens)
    return m_r(t, "p")


def m_var(t):          # italic variable (Greek letters, single letters)
    return m_r(t)


def _sub(base, sub):
    return "<m:sSub><m:e>%s</m:e><m:sub>%s</m:sub></m:sSub>" % (base, sub)


def _sSubSup(base, sub, sup):
    return "<m:sSubSup><m:e>%s</m:e><m:sub>%s</m:sub><m:sup>%s</m:sup></m:sSubSup>" % (base, sub, sup)


def _eq1():
    """RQ1 negative binomial: log(mu_t) = log(Pop_t) + b0 + b1(Year_t - Year_c), IRR = exp(b1)."""
    return (
        m_txt("log") + m_txt("(") + _sub(m_var("\u03bc"), m_var("t")) + m_txt(")")
        + m_txt("=")
        + m_txt("log") + m_txt("(") + _sub(m_txt("Pop"), m_var("t")) + m_txt(")")
        + m_txt("+")
        + _sub(m_var("\u03b2"), m_txt("0"))
        + m_txt("+")
        + _sub(m_var("\u03b2"), m_txt("1"))
        + m_txt("(") + _sub(m_txt("Year"), m_var("t")) + m_txt("\u2212") + _sub(m_txt("Year"), m_var("c")) + m_txt(")")
        + m_txt(", ")
        + m_txt("IRR") + m_txt("=") + m_txt("exp") + m_txt("(") + _sub(m_var("\u03b2"), m_txt("1")) + m_txt(")")
    )


def _eq2():
    """RQ2 logistic: logit[P(Fatal_i)] = b0 + b1 Sex_i + b2 Zone_i + b3 Role_i + b4 Year_i."""
    return (
        m_txt("logit") + m_txt("[") + m_txt("P") + m_txt("(") + _sub(m_txt("Fatal"), m_var("i")) + m_txt(")") + m_txt("]")
        + m_txt("=")
        + _sub(m_var("\u03b2"), m_txt("0"))
        + m_txt("+") + _sub(m_var("\u03b2"), m_txt("1")) + _sub(m_txt("Sex"), m_var("i"))
        + m_txt("+") + _sub(m_var("\u03b2"), m_txt("2")) + _sub(m_txt("Zone"), m_var("i"))
        + m_txt("+") + _sub(m_var("\u03b2"), m_txt("3")) + _sub(m_txt("Role"), m_var("i"))
        + m_txt("+") + _sub(m_var("\u03b2"), m_txt("4")) + _sub(m_txt("Year"), m_var("i"))
    )


def _eq3():
    """Bayesian mixed-effects: logit[P(Fatal_ij)] = X_ij b + u_j, u_j ~ N(0, sigma_u^2)."""
    return (
        m_txt("logit") + m_txt("[") + m_txt("P") + m_txt("(") + _sub(m_txt("Fatal"), m_var("ij")) + m_txt(")") + m_txt("]")
        + m_txt("=")
        + _sub(m_var("X"), m_var("ij")) + m_var("\u03b2")
        + m_txt("+") + _sub(m_var("u"), m_var("j"))
        + m_txt(", ")
        + _sub(m_var("u"), m_var("j")) + m_txt("\u223c") + m_txt("N")
        + m_txt("(") + m_txt("0") + m_txt(",")
        + _sSubSup(m_var("\u03c3"), m_var("u"), m_txt("2"))
        + m_txt(")")
    )


def add_display_equation(doc, inner):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    omath_para = parse_xml(
        "<m:oMathPara %s %s><m:oMath>%s</m:oMath></m:oMathPara>"
        % (NSDECL_M, NSDECL_W, inner)
    )
    p._p.append(omath_para)
    return p


# ------------------------------------------------------------ document formatting
doc = Document()

_normal = doc.styles["Normal"]
_normal.font.name = BODY_FONT
_normal.font.size = Pt(12)
_normal.paragraph_format.space_after = Pt(6)
_normal.paragraph_format.line_spacing = 1.15


def set_outline(p, level):
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml('<w:outlineLvl %s w:val="%d"/>' % (NSDECL_W, level)))


TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)", re.DOTALL)


def _add_run(p, text, bold=False, italic=False, size=None, code=False):
    text = text.replace("\\*", "*").replace("&nbsp;", "\u00a0")
    if not text:
        return
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    if code:
        r.font.name = "Consolas"
        r.font.size = Pt((size - 1) if size else 10)


def add_runs(p, text, base_bold=False, base_italic=False, base_size=None, base_code=False):
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            add_runs(p, part[2:-2], base_bold=True, base_italic=base_italic,
                     base_size=base_size, base_code=base_code)
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            _add_run(p, part[1:-1], bold=base_bold, italic=base_italic,
                     size=base_size, code=True)
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            add_runs(p, part[1:-1], base_bold=base_bold, base_italic=True,
                     base_size=base_size, base_code=base_code)
        else:
            _add_run(p, part, bold=base_bold, italic=base_italic,
                     size=base_size, code=base_code)


def add_para(doc, text="", size=None, bold=False, italic=False, align=None,
             space_after=6, line_spacing=1.15, left_indent=None,
             first_line_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    if text:
        add_runs(p, text, base_bold=bold, base_italic=italic, base_size=size)
    return p


def add_title(doc, text):
    return add_para(doc, text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=10, line_spacing=1.2)


def add_heading(doc, text, level):
    if level == 2:
        size, bold, italic, sb = 13, True, False, 14
        outline = 0
    else:  # level 3
        size, bold, italic, sb = 12, True, True, 10
        outline = 1
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    set_outline(p, outline)
    return p


def add_caption(doc, text):
    return add_para(doc, text, size=10, space_after=4)


# ------------------------------------------------------------------ tables
def parse_table_lines(lines):
    header = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    sep = lines[1].strip().strip("|")
    sepcells = [c.strip() for c in sep.split("|")]
    aligns = []
    for c in sepcells:
        if c.startswith(":") and c.endswith(":"):
            aligns.append("center")
        elif c.endswith(":"):
            aligns.append("right")
        else:
            aligns.append("left")
    rows = []
    for ln in lines[2:]:
        if ln.strip().startswith("|"):
            rows.append([c.strip() for c in ln.strip().strip("|").split("|")])
    return header, aligns, rows


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
}


def render_table(doc, header, aligns, rows):
    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    def fill(cell, text, bold, align):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        add_runs(p, text, base_bold=bold, base_size=10)

    for j, h in enumerate(header):
        fill(table.cell(0, j), h, True, _ALIGN_MAP.get(aligns[j], WD_ALIGN_PARAGRAPH.LEFT))
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            fill(table.cell(i + 1, j), val, False, _ALIGN_MAP.get(aligns[j], WD_ALIGN_PARAGRAPH.LEFT))

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    sp.paragraph_format.space_before = Pt(0)
    return table


# ------------------------------------------------------------------ images
def png_size(path):
    with open(path, "rb") as f:
        data = f.read(26)
    if data[:8] != b"\x89PNG\r\n\x1a\n":
        return None
    w, h = struct.unpack(">II", data[16:24])
    return w, h


def add_image(doc, relpath, md_dir):
    p = (md_dir / relpath).resolve()
    if not p.exists():
        p = (FIG_DIR / Path(relpath).name).resolve()
    if not p.exists():
        add_para(doc, "[missing figure: %s]" % relpath, size=10, italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    size = png_size(p)
    if size:
        w, h = size
        disp_w = min(6.0, 7.0 * (w / h))
    else:
        disp_w = 6.0
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    run.add_picture(str(p), width=Inches(disp_w))
    return para


# ------------------------------------------------------------------ main parse
lines = MD_PATH.read_text(encoding="utf-8").splitlines()
n = len(lines)
i = 0

EQ_ANCHORS = [
    ("reporting the incidence rate ratio (IRR).", _eq1),
    ("and year as covariates)", _eq2),
    ("not as a primary estimator.", _eq3),
]

in_frontmatter = False
in_references = False

while i < n:
    line = lines[i]
    stripped = line.strip()

    if stripped in ("---", "***", "___", "") :
        i += 1
        continue

    if stripped == "*Article*":
        i += 1
        continue

    if stripped.startswith("# "):
        add_title(doc, stripped[2:].strip())
        in_frontmatter = True
        i += 1
        continue

    if stripped.startswith("## "):
        heading = stripped[3:].strip()
        in_frontmatter = False
        if heading == "References":
            in_references = True
        add_heading(doc, heading, 2)
        i += 1
        continue

    if stripped.startswith("### "):
        add_heading(doc, stripped[4:].strip(), 3)
        i += 1
        continue

    if in_frontmatter:
        add_para(doc, stripped, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=4)
        i += 1
        continue

    if stripped.startswith("|"):
        j = i
        while j < n and lines[j].strip().startswith("|"):
            j += 1
        header, aligns, rows = parse_table_lines(lines[i:j])
        render_table(doc, header, aligns, rows)
        i = j
        continue

    if stripped.startswith("!["):
        m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
        if m:
            add_image(doc, m.group(1), MD_PATH.parent)
        i += 1
        continue

    if in_references and re.match(r"^\d+\.\s", stripped):
        add_para(doc, stripped, size=11, left_indent=0.5,
                 first_line_indent=-0.5, space_after=4)
        i += 1
        continue

    if (stripped.startswith("\u2020") or stripped.startswith("*Note:*")
            or stripped.startswith("Interpretation:")):
        add_para(doc, stripped, size=10, space_after=6)
        i += 1
        continue

    cap_match = re.match(r"\*\*(Table \d+|Table S\d+|Figure \d+|Figure S\d+)\.\*\*", stripped)
    if cap_match:
        label = cap_match.group(1)
        text = stripped
        if label in SIG_TABLES:
            text = text + SIG_LEGEND
        add_caption(doc, text)
        i += 1
        continue

    add_para(doc, stripped, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for anchor, eq_builder in EQ_ANCHORS:
        if anchor in stripped:
            try:
                add_display_equation(doc, eq_builder())
            except Exception as exc:  # noqa: BLE001
                print("WARNING: equation insert failed for anchor %r: %s" % (anchor, exc))
    i += 1

doc.save(str(OUT_PATH))
print("Saved:", OUT_PATH)

# ------------------------------------------------------------------ validation
d2 = Document(str(OUT_PATH))
print("paragraphs:", len(d2.paragraphs))
print("tables:", len(d2.tables))
print("images:", len(d2.inline_shapes))
xml = d2.element.xml
print("oMath count:", xml.count("<m:oMath>"))
print("significance legend occurrences:", xml.count("asterisk notation"))



