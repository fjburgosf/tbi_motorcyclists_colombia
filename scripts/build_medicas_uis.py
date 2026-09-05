# -*- coding: utf-8 -*-
"""
Build the Medicas UIS submission package (original article):
  * main manuscript .docx (English + Spanish resumen/palabras clave, Vancouver)
  * supplementary material .docx (Tables S1-S6, Figures S1-S2)
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

BASE = Path(__file__).resolve().parents[1]
MD = BASE / "manuscript" / "manuscript_medicas_uis_EN.md"
FIG = BASE / "figures"
OUTDIR = BASE / "manuscript" / "A someter en medicas uis"
OUTDIR.mkdir(parents=True, exist_ok=True)
MAIN_OUT = OUTDIR / "Manuscrito_Medicas_UIS.docx"
SUPP_OUT = OUTDIR / "Material_suplementario_Medicas_UIS.docx"

BODY_FONT = "Times New Roman"

RESUMEN = (
    "Introducción: los motociclistas concentran una proporción desproporcionada de las "
    "muertes por siniestros viales en América Latina y el traumatismo craneoencefálico es un "
    "determinante principal de esa mortalidad. No obstante, Colombia carecía de una "
    "caracterización nacional reproducible de esta problemática. Objetivo: caracterizar la "
    "mortalidad por traumatismo craneoencefálico en motociclistas entre 2015 y 2024 y evaluar "
    "la comparabilidad de las principales fuentes de datos del país. Metodología: se analizaron "
    "dos fuentes abiertas de microdatos, los registros forenses de lesiones por evento de "
    "transporte y las estadísticas vitales oficiales, vinculadas ecológicamente por departamento "
    "y año. La tendencia nacional se modeló con regresión binomial negativa y el desenlace fatal "
    "frente al no fatal con regresión logística multivariable y multinivel bayesiana. Resultados: "
    "el traumatismo craneoencefálico representó entre el 30 % y el 36 % de las muertes de "
    "motociclistas. La mortalidad específica aumentó a razón de 1,035 veces por año, concentrada "
    "entre 2022 y 2024. La desigualdad territorial fue marcada, con Casanare y Arauca como los "
    "departamentos de mayor tasa. Las víctimas fueron mayoritariamente hombres jóvenes "
    "conductores, y el registro vital subestimó las muertes entre 10 % y 25 % antes de 2022, "
    "convergiendo después de una integración del registro civil. Discusión y conclusiones: se "
    "aporta una caracterización nacional y reproducible de esta mortalidad y una "
    "advertencia cuantificada sobre la discontinuidad de cobertura del registro vital, relevante "
    "para la investigación con datos administrativos."
)
# Note: original-article resumen must be a single analytical paragraph <=250 words.

DECS_KEYWORDS = (
    "Traumatismos Craneoencefálicos; Accidentes de Tránsito; Motocicletas; Epidemiología; "
    "Medicina Legal; Estadísticas Vitales; Colombia; América Latina."
)
MESH_KEYWORDS = (
    "Brain Injuries, Traumatic; Accidents, Traffic; Motorcycles; Epidemiology; Forensic "
    "Medicine; Vital Statistics; Colombia; Latin America."
)

# ------------------------------------------------------------------ helpers
CITE_TOKEN = re.compile(r"(\[\d[\d,\s\u2013\u2014-]*\])")
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)", re.DOTALL)


def new_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    return doc


def _add_run(p, text, bold=False, italic=False, size=None, code=False):
    text = text.replace("&nbsp;", "\u00a0")
    if not text:
        return
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    if code:
        r.font.name = "Consolas"
        r.font.size = Pt((size - 1) if size else 10)


def _add_text(p, text, bold, italic, size, code, cites):
    if not cites:
        _add_run(p, text, bold=bold, italic=italic, size=size, code=code)
        return
    for seg in CITE_TOKEN.split(text):
        if not seg:
            continue
        if CITE_TOKEN.fullmatch(seg):
            inner = seg[1:-1].replace("\u2013", "-").replace("\u2014", "-")
            r = p.add_run(inner)
            r.font.superscript = True
            r.bold = bold
            r.italic = italic
            if size is not None:
                r.font.size = Pt(size)
        else:
            _add_run(p, seg, bold=bold, italic=italic, size=size, code=code)


def add_runs(p, text, base_bold=False, base_italic=False, base_size=None,
             base_code=False, cites=False):
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            add_runs(p, part[2:-2], True, base_italic, base_size, base_code, cites)
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            _add_run(p, part[1:-1], bold=base_bold, italic=base_italic, size=base_size, code=True)
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            add_runs(p, part[1:-1], base_bold, True, base_size, base_code, cites)
        else:
            _add_text(p, part, base_bold, base_italic, base_size, base_code, cites)


def add_para(doc, text="", size=None, bold=False, italic=False, align=None,
             cites=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        add_runs(p, text, base_bold=bold, base_italic=italic, base_size=size, cites=cites)
    return p


# ------------------------------------------------------------------ Vancouver references
def _vancouver_authors(a):
    tokens = [t.strip() for t in a.split("; ") if t.strip()]
    out = []
    for t in tokens:
        if t == "et al.":
            out.append("et al.")
            continue
        t2 = re.sub(r"(?<=[A-Z\u00c0-\u00dd])\.", "", t)  # drop periods after initials
        t2 = t2.rstrip(".")
        t2 = t2.replace(", Jr", " Jr").replace(" Jr.", " Jr")
        out.append(t2)
    return ", ".join(out)


def _split_author_title(head):
    if "et al." in head:
        p = head.index("et al.") + len("et al.")
        return head[:p].strip(), head[p:].strip().lstrip(".").strip()
    if "; " in head:
        parts = head.split("; ")
        last = parts[-1]
        m = re.match(r"^(.*?[A-Z\u00c0-\u00dd]\.(?:[A-Z\u00c0-\u00dd]\.)*(?:\s*,?\s*Jr\.?)?)\s+(.+)$", last)
        if m:
            return "; ".join(parts[:-1] + [m.group(1).strip()]), m.group(2).strip()
    else:
        m = re.match(r"^(.*?[A-Z\u00c0-\u00dd]\.(?:[A-Z\u00c0-\u00dd]\.)*(?:\s*,?\s*Jr\.?)?)\s+(.+)$", head)
        if m:
            return m.group(1).strip(), m.group(2).strip()
    return "", head


def vancouver_ref(raw):
    s = re.sub(r"^\d+\.\s+", "", raw.strip())
    doi = ""
    m = re.search(r"https://doi\.org/(\S+?)\s*$", s)
    if m:
        doi = m.group(1).rstrip(".")
        s = s[:m.start()].strip()
    journal = year = vol = pages = ""
    jm = re.search(r"\*(.+?)\*", s)
    if jm:
        journal = jm.group(1).strip()
        head = s[:jm.start()].strip()
        tail = s[jm.end():].strip()
        tm = re.match(r"\*\*(\d{4})\*\*,\s*\*(.+?)\*,\s*(.+?)\s*$", tail)
        if tm:
            year, vol, pages = tm.group(1), tm.group(2), tm.group(3).rstrip(".")
        else:
            tm2 = re.match(r"\*\*(\d{4})\*\*", tail)
            if tm2:
                year = tm2.group(1)
    else:
        head = s.strip()
    authors, title = _split_author_title(head)
    va = _vancouver_authors(authors) if authors else ""
    out = (va.rstrip(".") + ". " if va else "") + title.rstrip(".") + "."
    if journal:
        out += " " + journal.rstrip(".") + "."
    if year:
        out += " " + year
        if vol:
            out += ";" + vol
            if pages:
                out += ":" + pages.replace("\u2013", "-").replace("\u2014", "-")
        out += "."
    if doi:
        out += " doi:" + doi
    return out


# ------------------------------------------------------------------ tables / images
def _png_size(path):
    import struct
    with open(path, "rb") as f:
        d = f.read(26)
    if d[:8] != b"\x89PNG\r\n\x1a\n":
        return None
    return struct.unpack(">II", d[16:24])


_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
          "center": WD_ALIGN_PARAGRAPH.CENTER}


def parse_table(lines):
    header = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    sepcells = [c.strip() for c in lines[1].strip().strip("|").split("|")]
    aligns = []
    for c in sepcells:
        if c.startswith(":") and c.endswith(":"):
            aligns.append("center")
        elif c.endswith(":"):
            aligns.append("right")
        else:
            aligns.append("left")
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")]
            for ln in lines[2:] if ln.strip().startswith("|")]
    return header, aligns, rows


def render_table(doc, header, aligns, rows):
    ncols = len(header)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    def fill(cell, text, bold, align):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        add_runs(p, text, base_bold=bold, base_size=10)

    for j, h in enumerate(header):
        fill(tbl.cell(0, j), h, True, aligns[j])
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            fill(tbl.cell(i + 1, j), val, False, aligns[j])
    return tbl


def add_image(doc, relpath):
    p = (MD.parent / relpath).resolve()
    if not p.exists():
        p = (FIG / Path(relpath).name).resolve()
    if not p.exists():
        add_para(doc, "[figura no encontrada: %s]" % relpath, size=10, italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    sz = _png_size(p)
    disp_w = min(5.5, 7.0 * (sz[0] / sz[1])) if sz else 5.5
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    para.add_run().add_picture(str(p), width=Inches(disp_w))
    return para


# ------------------------------------------------------------------ manuscript build
def _strip_num(t):
    return re.sub(r"^\d+(\.\d+)*\.\s*", "", t)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(_strip_num(text))
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(_strip_num(text))
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    return p


ETHICS_TEXT = (
    "Ethical considerations. This study used exclusively publicly available, deidentified "
    "secondary data (individual level microdata published by the National Institute of Legal "
    "Medicine and Forensic Sciences and by the National Administrative Department of Statistics "
    "through the Colombian open data portal, together with official aggregated population "
    "statistics), none of which can be traced to identifiable individuals. No new data were "
    "collected from human participants and no identifiable personal information was accessed. "
    "Therefore, approval by a research ethics committee was not required, in accordance with "
    "Resolution 8430 of 1993 of the Colombian Ministry of Health."
)


def process_body(doc, lines):
    i = 0
    n = len(lines)
    ethics_added = False
    while i < n:
        line = lines[i]
        s = line.strip()
        if s in ("---", "", "***", "___"):
            i += 1
            continue
        if s.startswith("### "):
            add_h2(doc, s[4:].strip())
            i += 1
            continue
        if s.startswith("## "):
            if s.startswith("## 3.") and not ethics_added:
                add_h2(doc, "Ethical considerations")
                add_para(doc, ETHICS_TEXT, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                ethics_added = True
            add_h1(doc, s[3:].strip())
            i += 1
            continue
        if s.startswith("|"):
            j = i
            while j < n and lines[j].strip().startswith("|"):
                j += 1
            header, aligns, rows = parse_table(lines[i:j])
            render_table(doc, header, aligns, rows)
            i = j
            continue
        if s.startswith("!["):
            m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", s)
            if m:
                add_image(doc, m.group(1))
            i += 1
            continue
        if re.match(r"\*\*(Table|Figure)\s+\d+\.\*\*", s):
            add_para(doc, s + " Source: own elaboration.", size=10,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
            i += 1
            continue
        if s.startswith("\u2020") or s.startswith("*Note"):
            add_para(doc, s, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            i += 1
            continue
        add_para(doc, s, align=WD_ALIGN_PARAGRAPH.JUSTIFY, cites=True)
        i += 1
    return ethics_added


def build_main():
    doc = new_document()
    text = MD.read_text(encoding="utf-8")

    intro = text.find("## 1. Introduction")
    supp = text.find("## Supplementary Materials")
    refs_idx = text.find("## References")

    # title
    title = ""
    for ln in text[:intro].splitlines():
        if ln.strip().startswith("# "):
            title = ln.strip()[2:].strip()
            break
    add_para(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # authors / credits / correspondence
    add_para(doc, "Francisco Burgos-Florez\u00b9", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    add_para(doc,
             "\u00b9 Escuela de Pregrado, Direcci\u00f3n Acad\u00e9mica, Vicerrector\u00eda de "
             "Sede, Universidad Nacional de Colombia, Sede La Paz. Cesar. Colombia. "
             "ORCID: [0000-0002-5381-1398]. CvLAC: https://scienti.minciencias.gov.co/"
             "cvlac/visualizador/generarCurriculoCv.do?cod_rh=0001615865",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    add_para(doc,
             "Correspondencia: Francisco Burgos-Florez. Correo electr\u00f3nico: "
             "fjburgosf@unal.edu.co",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

    # resumen (ES)
    add_h1(doc, "Resumen")
    add_para(doc, RESUMEN, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "Palabras clave: " + DECS_KEYWORDS, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=12)

    # abstract (EN)
    abstract_par = ""
    lines_before = text[:intro].splitlines()
    for k, ln in enumerate(lines_before):
        if ln.strip() == "## Abstract":
            # take the first non-empty line after the heading (skip blanks)
            for m in range(k + 1, len(lines_before)):
                if lines_before[m].strip():
                    abstract_par = lines_before[m].strip()
                    break
            break
    add_h1(doc, "Abstract")
    if abstract_par:
        add_para(doc, abstract_par, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "Keywords: " + MESH_KEYWORDS, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

    # body
    process_body(doc, text[intro:supp].splitlines())

    # funding / conflicts / contributions
    add_h1(doc, "Funding sources")
    add_para(doc, "This research received no external funding.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_h1(doc, "Conflicts of interest")
    add_para(doc, "The author declares no conflicts of interest.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_h1(doc, "Author contributions")
    add_para(doc,
             "Francisco Burgos-Florez: conceptualization, methodology, software, validation, "
             "formal analysis, investigation, data curation, writing\u2014original draft, "
             "writing\u2014review and editing, visualization, project administration.",
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # material suplementario note
    add_h1(doc, "Material suplementario")
    add_para(doc,
             "Tablas S1\u2013S6 y Figuras S1\u2013S2 se presentan en un archivo aparte "
             "(Material_suplementario_Medicas_UIS.docx).",
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # references (Vancouver)
    add_h1(doc, "References")
    ref_lines = text[refs_idx:].splitlines()
    nrefs = 0
    for ln in ref_lines:
        s = ln.strip()
        if re.match(r"^\d+\.\s", s):
            add_para(doc, vancouver_ref(s), size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_after=4)
            nrefs += 1

    doc.save(str(MAIN_OUT))
    return nrefs


def build_supp():
    doc = new_document()
    text = MD.read_text(encoding="utf-8")
    supp = text.find("## Supplementary Materials")
    back = text.find("## Back Matter")
    lines = text[supp:back].splitlines()

    add_para(doc, "Material suplementario", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    i = 0
    n = len(lines)
    ntables = 0
    nfigs = 0
    while i < n:
        s = lines[i].strip()
        if s in ("---", "", "***", "___"):
            i += 1
            continue
        if s.startswith("## "):
            i += 1
            continue
        if "reproducible pipeline" in s:
            i += 1
            continue
        if s.startswith("|"):
            j = i
            while j < n and lines[j].strip().startswith("|"):
                j += 1
            header, aligns, rows = parse_table(lines[i:j])
            render_table(doc, header, aligns, rows)
            ntables += 1
            i = j
            continue
        if s.startswith("!["):
            m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", s)
            if m:
                add_image(doc, m.group(1))
                nfigs += 1
            i += 1
            continue
        if re.match(r"\*\*(Table|Figure)\s+[S]?\d+\.\*\*", s):
            add_para(doc, s + " Source: own elaboration.", size=10,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
            i += 1
            continue
        if s.startswith("\u2020") or s.startswith("*Note"):
            add_para(doc, s, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            i += 1
            continue
        add_para(doc, s, align=WD_ALIGN_PARAGRAPH.JUSTIFY, cites=False)
        i += 1

    doc.save(str(SUPP_OUT))
    return ntables, nfigs


if __name__ == "__main__":
    nrefs = build_main()
    ntables, nfigs = build_supp()
    print("Saved main manuscript:", MAIN_OUT)
    print("  references (Vancouver):", nrefs)
    print("Saved supplementary:", SUPP_OUT)
    print("  supplementary tables:", ntables, "| supplementary figures:", nfigs)

    # quick validation
    d1 = Document(str(MAIN_OUT))
    d2 = Document(str(SUPP_OUT))
    print("\nMAIN  -> paragraphs:", len(d1.paragraphs), "tables:", len(d1.tables),
          "images:", len(d1.inline_shapes))
    print("SUPP  -> paragraphs:", len(d2.paragraphs), "tables:", len(d2.tables),
          "images:", len(d2.inline_shapes))
    # count superscript citations in main
    sup = 0
    for p in d1.paragraphs:
        for r in p.runs:
            if r.font.superscript:
                sup += 1
    print("superscript citation runs in main:", sup)




