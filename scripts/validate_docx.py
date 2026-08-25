# -*- coding: utf-8 -*-
"""Validate the generated docx: headings, table headers, images, equations, captions."""
from docx import Document
from docx.oxml.ns import qn

P = r"D:/ACADEMICO/Papers/TBI_data_analysis/manuscript/manuscript_safety_final_EN.docx"
d = Document(P)

print("=== HEADINGS / bold short paragraphs ===")
for p in d.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    if t.startswith(("1.", "2.", "3.", "4.", "5.")) or t in ("Abstract", "Supplementary Materials", "Back Matter", "References"):
        print("H:", t)

print("\n=== TABLE HEADERS ===")
for ti, tbl in enumerate(d.tables):
    hdr = [c.text for c in tbl.rows[0].cells]
    print(f"Table {ti+1}: {hdr}")

print("\n=== IMAGES ===")
for shp in d.inline_shapes:
    print(shp.width, shp.height, shp.type)

print("\n=== EQUATIONS (oMath text) ===")
for p in d.paragraphs:
    omaths = p._p.findall(qn('m:oMath'))
    for om in omaths:
        texts = [t.text or '' for t in om.iter(qn('m:t'))]
        print(''.join(texts))

print("\n=== CAPTIONS WITH SIGNIFICANCE LEGEND ===")
for p in d.paragraphs:
    if "asterisk notation" in p.text:
        print(p.text[:110], "...")

print("\n=== FIGURE/TABLE CAPTIONS ===")
for p in d.paragraphs:
    t = p.text.strip()
    if t.startswith(("Figure ", "Table ")):
        print(t[:100])
